import subprocess
import sys

def install(package):
    """Install the specified Python package using pip."""
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

def check_and_install_packages(required_packages):
    missing_packages = []
    installed_packages = []  # Keep track of newly installed packages

    for package in required_packages:
        try:
            __import__(package)
            print(f"{package} is already installed.")
        except ImportError:
            missing_packages.append(package)

    if missing_packages:
        missing_packages_str = ", ".join(missing_packages)
        consent = input(f"Would you like to install the missing packages ({missing_packages_str})? [yes/no]: ")
        if consent.lower() == 'yes':
            for package in missing_packages:
                print(f"Installing {package}...")
                install(package)
                installed_packages.append(package)  # Add to the list of installed packages
        else:
            print("Missing packages are required to run the application. Exiting.")
            sys.exit(1)
    else:
        print("All required packages are already installed.")

    # Summarising Installed Packages
    if installed_packages:
        print(f"Installed packages: {', '.join(installed_packages)}")
    else:
        print("No new packages were installed.")

    # Wait for User Input
    input("Press Enter to continue with the script...")

# List of required packages
required_packages = ["python-docx", "os"]  # Note: corrected package name to lowercase for consistency

# Check and install required packages
check_and_install_packages(required_packages)

# main script - Split files into seperate documents
from docx import Document
import os

def get_document_path():
    return input("Please enter the path to the original document: ")

def get_destination_path():
    return input("Please enter the path for saving extracted chapters: ")

def get_heading_level():
    return input("Please select the chapter heading level (for <h1> type 1, for <h2> type 2, etc.): ")

def validate_heading_selection(doc, heading_style):
    # Count the number of headings matching the selected style
    count = sum(1 for paragraph in doc.paragraphs if paragraph.style.name == heading_style)
    print(f"Number of headers found with the style '{heading_style}': {count}")
    if count == 0:
        print("Warning: No headers found with the selected style. Please reconsider your selection.")
    return count

def user_confirmation():
    return input("Would you like to proceed with the extraction? Type 'yes' to continue or 'no' to select a different heading level: ").lower()

def save_chapter(chapter_title, chapter_content, destination_path):
    if not os.path.exists(destination_path):
        os.makedirs(destination_path)
    chapter_doc = Document()
    chapter_doc.add_heading(chapter_title, level=1)
    for paragraph in chapter_content:
        chapter_doc.add_paragraph(paragraph)
    filename = f"{chapter_title}.docx".replace(":", "").replace("?", "").replace("/", "").replace("\\", "")
    chapter_doc.save(os.path.join(destination_path, filename))

def extract_chapters(doc, heading_style, destination_path):
    current_chapter_content = []
    current_chapter_title = ""
    for paragraph in doc.paragraphs:
        if paragraph.style.name == heading_style:
            if current_chapter_content:
                save_chapter(current_chapter_title, current_chapter_content, destination_path)
                current_chapter_content = []
            current_chapter_title = paragraph.text
        else:
            current_chapter_content.append(paragraph.text)
    if current_chapter_content:
        save_chapter(current_chapter_title, current_chapter_content, destination_path)
    print("All chapters have been extracted and saved as individual docx files.")

def main():
    # Main process
    doc_path = get_document_path()
    destination_path = get_destination_path()
    doc = Document(doc_path)

    while True:
        heading_level = get_heading_level()
        heading_style = f"Heading {heading_level}"
        if validate_heading_selection(doc, heading_style) > 0:
            if user_confirmation() == "yes":
                extract_chapters(doc, heading_style, destination_path)
                break
        else:
            if user_confirmation() == "no":
                continue
            else:
                extract_chapters(doc, heading_style, destination_path)
                break

if __name__ == "__main__":
    main()